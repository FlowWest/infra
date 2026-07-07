"""
generate_report.py — FlowWest Monthly Progress Report generator

Usage:
    python generate_report.py <input.json>

Reads a JSON file describing the report contents and writes a .docx to the
path specified by output_path in that file.

Relative paths in the JSON are resolved against the skill root (the directory
containing this scripts/ folder). So "references/flowwest_logo.jpg" always
resolves correctly regardless of where you call the script from.

Input JSON schema:
{
  "output_path": "../../June_2026_028-11_Progress_Report.docx",  // relative or absolute
  "logo_path": "references/flowwest_logo.jpg",   // optional — skip logo if missing
  "sig_path": "/abs/path/to/signature.png",       // optional — omit for no signature
  "recipient": {
    "name": "Kim Brewitt",
    "title": "Task Order Manager",
    "org": "California Department of Water Resources",
    "address": "1416 9th Street",
    "city_state_zip": "Sacramento, CA 95814",
    "phone": "(916) 651-9370",
    "email": "kim.brewitt@water.ca.gov"
  },
  "contract": {
    "contract_no": "4600014515",
    "to_name": "Task Order 028-11: CDEC Sonde Tracker",
    "project_no": "028-11"
  },
  "date": "July 6, 2026",
  "intro": "Enclosed please find FlowWest's invoice for work...",
  "deliverable": null,    // or a string — inserted as a second paragraph after intro
  "tasks": [
    { "num": "1", "name": "Project Management", "lines": ["bullet 1", "bullet 2"] },
    { "num": "2", "name": "App Development",    "lines": [] }  // empty = no work completed
  ],
  "remaining_balance": "342,155.00",
  "pm_name": "Sadie Gill",
  "pm_title": "Principal Data Scientist"
}
"""

import json
import os
import sys

from docx import Document
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH


# ── helpers ───────────────────────────────────────────────────────────────────

def no_space(p):
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def remove_table_borders(table):
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'none')
        tblBorders.append(el)
    tblPr.append(tblBorders)


def set_table_width(table, width_twips):
    """Pin table to an explicit width with fixed layout (prevents margin bleed)."""
    tbl   = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(width_twips))
    tblW.set(qn('w:type'), 'dxa')
    tblPr.append(tblW)
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)


def set_table_col_widths(table, widths_twips):
    """Set column widths via tblGrid + each cell. Call AFTER all rows are added."""
    tbl     = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(0, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in widths_twips:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(w))
        tblGrid.append(gc)
    for row in table.rows:
        seen = set()
        col  = 0
        for cell in row.cells:
            if id(cell) in seen:
                continue
            seen.add(id(cell))
            if col >= len(widths_twips):
                break
            tc   = cell._tc
            tcPr = tc.find(qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                tc.insert(0, tcPr)
            tcW = tcPr.find(qn('w:tcW'))
            if tcW is None:
                tcW = OxmlElement('w:tcW')
                tcPr.append(tcW)
            tcW.set(qn('w:w'),    str(widths_twips[col]))
            tcW.set(qn('w:type'), 'dxa')
            col += 1


def zero_cell_right_margin(cell):
    """Remove right padding from a table cell."""
    tc   = cell._tc
    tcPr = tc.find(qn('w:tcPr'))
    if tcPr is None:
        tcPr = OxmlElement('w:tcPr')
        tc.insert(0, tcPr)
    tcMar = tcPr.find(qn('w:tcMar'))
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    right = tcMar.find(qn('w:right'))
    if right is None:
        right = OxmlElement('w:right')
        tcMar.append(right)
    right.set(qn('w:w'),    '0')
    right.set(qn('w:type'), 'dxa')


def add_horizontal_rule(doc):
    p   = doc.add_paragraph()
    no_space(p)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── main ──────────────────────────────────────────────────────────────────────

def generate(data: dict):
    logo_path = data.get('logo_path', '')
    sig_path  = data.get('sig_path', '')
    r         = data['recipient']
    c         = data['contract']

    doc = Document()
    section = doc.sections[0]
    section.page_width    = Inches(8.5)
    section.page_height   = Inches(11)
    section.left_margin   = Inches(1)
    section.right_margin  = Inches(1)
    section.top_margin    = Inches(1)
    section.bottom_margin = Inches(1)

    # ── 1. Header table: logo + address ───────────────────────────────────────
    HCOL_LEFT  = int(3.5 * 1440)   # 5040 twips
    HCOL_RIGHT = int(3.0 * 1440)   # 4320 twips

    htable = doc.add_table(rows=1, cols=2)
    remove_table_borders(htable)
    set_table_width(htable, HCOL_LEFT + HCOL_RIGHT)

    right_cell = htable.rows[0].cells[1]
    zero_cell_right_margin(right_cell)

    if logo_path and os.path.exists(logo_path):
        lp = right_cell.paragraphs[0]
        lp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        no_space(lp)
        lp.add_run().add_picture(logo_path, width=Inches(1.75))

    ap = right_cell.add_paragraph("P.O. Box 29392\nOakland, CA 94604") \
         if not (logo_path and os.path.exists(logo_path)) \
         else right_cell.add_paragraph("P.O. Box 29392\nOakland, CA 94604")
    ap.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    no_space(ap)
    for run in ap.runs:
        run.font.size = Pt(9)

    set_table_col_widths(htable, [HCOL_LEFT, HCOL_RIGHT])

    # ── 2. Title + rule ────────────────────────────────────────────────────────
    tp = doc.add_paragraph()
    no_space(tp)
    tp.paragraph_format.space_before = Pt(6)
    tr = tp.add_run("Monthly Progress Report")
    tr.bold = True
    tr.font.size = Pt(13)

    rule_p = add_horizontal_rule(doc)
    rule_p.paragraph_format.space_before = Pt(2)
    rule_p.paragraph_format.space_after  = Pt(8)
    doc.add_paragraph()   # blank line after rule

    # ── 3. Recipient block ─────────────────────────────────────────────────────
    for line in [r['name'], r['title'], r['org'], r['address'], r['city_state_zip'], r['phone'], r['email']]:
        if line and line.strip():
            p = doc.add_paragraph(line)
            no_space(p)

    doc.add_paragraph()   # blank line

    for line in [f"Contract No.: {c['contract_no']}", c['to_name'], f"(FlowWest Project No.: {c['project_no']})"]:
        p = doc.add_paragraph(line)
        no_space(p)

    # ── 4. Date ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    dp = doc.add_paragraph(data['date'])
    no_space(dp)

    # ── 5. Salutation ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    sp = doc.add_paragraph(f"Dear {r['name']},")
    no_space(sp)

    # ── 6. Body ────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph(data['intro'])

    if data.get('deliverable'):
        doc.add_paragraph(data['deliverable'])

    doc.add_paragraph("Work completed during this billing period is summarized below.")

    # ── 7. Task table ──────────────────────────────────────────────────────────
    TASK_TW  = 2340    # 1.625" — 1/4 of 6.5"
    DESC_TW  = 7020    # 4.875" — 3/4 of 6.5"
    TOTAL_TW = TASK_TW + DESC_TW   # 9360 twips = 6.5"

    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'

    hdr = table.rows[0].cells
    for cell in hdr:
        no_space(cell.paragraphs[0])
    hdr[0].paragraphs[0].clear()
    hdr[0].paragraphs[0].add_run("Task").bold = True
    hdr[1].paragraphs[0].clear()
    hdr[1].paragraphs[0].add_run("Description").bold = True

    for task in data['tasks']:
        row = table.add_row().cells

        p0 = row[0].paragraphs[0]
        p0.clear()
        no_space(p0)
        p0.add_run(str(task['num'])).bold = True

        p1 = row[1].paragraphs[0]
        p1.clear()
        no_space(p1)
        p1.add_run(task['name']).bold = True

        lines = task.get('lines', [])
        if lines:
            for line in lines:
                p1.add_run('\n• ' + line)
        else:
            p1.add_run('\n')
            p1.add_run('No work completed under this task.').italic = True

    # Budget and Schedule — merged row
    brow = table.add_row().cells
    brow[0].merge(brow[1])
    bp = brow[0].paragraphs[0]
    bp.clear()
    no_space(bp)
    bp.add_run('Budget and Schedule').bold = True
    bp.add_run('\n')
    bp.add_run(f"Remaining balance is ${data['remaining_balance']}").italic = True

    set_table_col_widths(table, [TASK_TW, DESC_TW])
    set_table_width(table, TOTAL_TW)

    # ── 8. Closing ─────────────────────────────────────────────────────────────
    doc.add_paragraph()
    doc.add_paragraph(
        "Thank you for the continuing opportunity to work on this project. "
        "If you have any questions or comments regarding these charges, "
        "please do not hesitate to contact me."
    )

    doc.add_paragraph("Sincerely,")

    if sig_path and os.path.exists(sig_path):
        doc.add_paragraph()
        sigp = doc.add_paragraph()
        sigp.add_run().add_picture(sig_path, width=Inches(2))
        doc.add_paragraph()

    p_name = doc.add_paragraph(data['pm_name']);  no_space(p_name)
    p_title = doc.add_paragraph(data['pm_title']); no_space(p_title)
    p_org = doc.add_paragraph("FlowWest, LLC");   no_space(p_org)

    # ── save ───────────────────────────────────────────────────────────────────
    out = data['output_path']
    os.makedirs(os.path.dirname(out), exist_ok=True) if os.path.dirname(out) else None
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == '__main__':
    if len(sys.argv) != 2:
        print("Usage: python generate_report.py <input.json>")
        sys.exit(1)
    with open(sys.argv[1]) as f:
        data = json.load(f)

    # Resolve relative paths against the skill root (parent of scripts/)
    skill_root = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    for key in ('output_path', 'logo_path', 'sig_path'):
        val = data.get(key)
        if val and not os.path.isabs(val):
            data[key] = os.path.join(skill_root, val)

    generate(data)
